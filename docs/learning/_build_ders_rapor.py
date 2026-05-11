"""Build TurkResearcher_Ders_Rapor.docx — pedagogical/tutorial-style learning doc.

Açıklar: kullanılan tüm terimler, işlemler, kavramlar; proje üzerinden anlatılır.
Hedef: bir öğrenci/junior bu dokümanla projeyi ve arkasındaki teorinin çoğunu
kavrayabilsin.

Run:
    python docs/learning/_build_ders_rapor.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(__file__).resolve().parent
OUT = HERE / "TurkResearcher_Ders_Rapor.docx"


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_heading(doc, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        if level == 0:
            run.font.size = Pt(20)
        elif level == 1:
            run.font.size = Pt(14)
        elif level == 2:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)


def add_para(doc, text: str, *, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_bullet(doc, text: str, level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)


def add_term(doc, term: str, definition: str, *, in_project: str = "") -> None:
    """A term entry: bold name, definition, then how it shows up in this project."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{term}. ")
    r1.bold = True; r1.font.name = "Times New Roman"; r1.font.size = Pt(11)
    r2 = p.add_run(definition)
    r2.font.name = "Times New Roman"; r2.font.size = Pt(11)

    if in_project:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.left_indent = Inches(0.25)
        r1 = p2.add_run("Bu projede: ")
        r1.italic = True; r1.bold = True; r1.font.name = "Times New Roman"; r1.font.size = Pt(10)
        r2 = p2.add_run(in_project)
        r2.italic = True; r2.font.name = "Times New Roman"; r2.font.size = Pt(10)


def add_meta_row(doc, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    r1.font.name = "Times New Roman"; r1.font.size = Pt(11); r1.bold = True
    r2 = p.add_run(value)
    r2.font.name = "Times New Roman"; r2.font.size = Pt(11)


def add_table(doc, headers, rows,
              header_fill="1F4E79", header_color=RGBColor(0xFF, 0xFF, 0xFF),
              col_widths=None) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(10.5)
        run.font.color.rgb = header_color
        set_cell_shading(hdr[i], header_fill)
    for ri, row in enumerate(rows, start=1):
        cells = tbl.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Times New Roman"; r.font.size = Pt(10)
    if col_widths:
        for row in tbl.rows:
            for c, w in zip(row.cells, col_widths):
                c.width = Inches(w)
    doc.add_paragraph()


def build():
    doc = Document()

    # ─── Cover ───
    add_heading(doc, "TürkResearcher", level=0)
    add_heading(doc, "Ders Niyetine Teknik Rehber", level=1)

    add_meta_row(doc, "Kim için", "Bu projeyi anlamak veya benzer bir sistem kuracak öğrenci/junior")
    add_meta_row(doc, "Format", "Kavram → Tanım → Bu projedeki kullanım — sırasıyla")
    add_meta_row(doc, "Süre", "Doğru yerde duraklayarak okuyana ~2-3 saatlik bir ders niteliğindedir")
    add_meta_row(doc, "Tarih", "Mayıs 2026")
    doc.add_paragraph()

    # ─── 0. Genel Çerçeve ───
    add_heading(doc, "0. Önce Büyük Resim", level=1)
    add_para(doc,
        "TürkResearcher; Türkçe akademik literatürde sorulan sorulara, gerçek tezlerden alıntı "
        "yapan, kendi kalitesini ölçen bir araştırma asistanıdır. Sade ifadeyle: "
        "kullanıcı Türkçe bir soru sorar, sistem 633 binden fazla tezden ilgili olanları bulur, "
        "bir LLM bu kaynakları okur ve atıflı bir özet yazar. \"İlgili olanı bulma\" işine "
        "Information Retrieval (IR), \"yazma\" işine Generation deriz; ikisi birleşince "
        "Retrieval-Augmented Generation (RAG) olur."
    )
    add_para(doc, "Üç ana taş üzerine kurulu:", bold=True)
    add_bullet(doc, "NLP — Türkçe metni anlamlandırma, vektörleştirme (embedding), arama.")
    add_bullet(doc, "LLM — DeepSeek modelini bir agent olarak (planner, writer, judge) kullanma.")
    add_bullet(doc, "MLOps — Veriyi, modeli, demoyu yeniden üretilebilir biçimde yayınlama.")

    # ─── 1. Temel Terimler ───
    add_heading(doc, "1. Temel Terimler (Bunları Bilmeden Geçilmez)", level=1)

    add_term(doc, "NLP (Doğal Dil İşleme)",
        "Bilgisayarın insan dilini işleyebilmesi için kullanılan teknik ve teorilerin "
        "tamamı. Tokenizasyondan dil modellerine kadar her şey bu şemsiyenin altındadır.",
        in_project="Türkçe abstract'ları işlemek, embedding üretmek, retrieval yapmak.")

    add_term(doc, "LLM (Büyük Dil Modeli)",
        "Çok büyük metin verisinde önceden eğitilmiş üretken model (örn. GPT, Claude, "
        "DeepSeek, Llama). Doğal dilde sorular cevaplar, metin yazar, kod üretir.",
        in_project="DeepSeek API; planner ve writer agent'larının arkasındaki model.")

    add_term(doc, "RAG (Retrieval-Augmented Generation)",
        "LLM'i \"hazır bilgi havuzuyla\" desteklemek. Soru gelir → ilgili dokümanlar "
        "vektör aramayla bulunur → LLM bu dokümanları görerek cevap üretir. Hallucination'ı "
        "azaltır, atıflı çıktı üretir.",
        in_project="Bu sistemin temel mimarisi. Tezler vektör DB'de, LLM onları okuyarak yazar.")

    add_term(doc, "Embedding",
        "Bir metni (genelde 384, 768 ya da 1024 boyutlu) sabit uzunluklu sayı vektörüne "
        "çevirme. Anlamca yakın metinler vektör uzayında yakın olacak şekilde eğitilir.",
        in_project="Her tez 768-dim vektöre çevrilir; sorgu da aynı uzayda aranır.")

    add_term(doc, "Cosine Similarity / Distance",
        "İki vektör arasındaki açının kosinüsü. 1 = aynı yön (en benzer), 0 = ortogonal, "
        "-1 = zıt. Cosine distance = 1 - cosine similarity. Embeddings normalize edildiğinde "
        "kosinüs benzerlik = nokta çarpım.",
        in_project="Chroma'da hnsw:space=cosine ile kullanılıyor; retrieval'da score = 1 - distance/2.")

    add_term(doc, "Tokenization",
        "Metni modelin anlayabileceği parçalara (token) bölme. Subword tokenizer'lar "
        "(SentencePiece, BPE, WordPiece) bilinmeyen kelimeleri parçalara ayırarak handle eder.",
        in_project="mpnet'in SentencePiece tokenizer'ı kullanılıyor; max_seq=256 token sınırı.")

    add_term(doc, "Transformer",
        "Self-attention mekanizmasına dayalı sinir ağı mimarisi (\"Attention Is All You "
        "Need\", 2017). Hem encoder (BERT-tipi) hem decoder (GPT-tipi) hem encoder-decoder "
        "(T5-tipi) varyasyonları var.",
        in_project="mpnet bir BERT/RoBERTa türevi encoder transformer; embedding üretir.")

    add_term(doc, "Pre-training vs Fine-tuning",
        "Pre-training: model çok büyük genel veride sıfırdan eğitilir (pahalı, GPU-saat). "
        "Fine-tuning: pre-trained model küçük domain verisinde ince ayar yapılır (ucuz, "
        "saatler). Bu projede sadece fine-tune yapıyoruz.",
        in_project="trakad-embed-v2 = mpnet üzerine 633K tezle yapılan kontrastif fine-tune.")

    add_term(doc, "Sentence-Transformers (sbert)",
        "Cümle/paragraf seviyesinde embedding üreten kütüphane. SentenceTransformer "
        "sınıfı, fit(), encode() gibi yüksek seviyeli API'ler sunar.",
        in_project="Hem v1 (mpnet) yüklemek hem v2'yi fine-tune etmek için kullanıldı.")

    # ─── 2. Veri Tarafı ───
    add_heading(doc, "2. Veri Pipeline'ı — Korpusu Nereden ve Nasıl Aldık", level=1)

    add_term(doc, "OAI-PMH",
        "Open Archives Initiative Protocol for Metadata Harvesting. Akademik kütüphanelerin "
        "(YÖK Tez dahil) metadata yayınladığı standart. HTTP üzerinden XML tabanlı sayfalı "
        "veri çeker (resumptionToken ile pagination).",
        in_project="YÖK Tez OAI-PMH endpoint'inden ~1M tez metadata'sı harvest edildi.")

    add_term(doc, "Dublin Core",
        "Metadata standartı; title, creator, date, description, subject gibi standardize "
        "alanlar tanımlar. OAI-PMH genelde Dublin Core formatında veri yayınlar.",
        in_project="YÖK Tez kayıtları Dublin Core olarak iniyor; parser bunu Python dict'e çeviriyor.")

    add_term(doc, "Apache Parquet",
        "Sütun bazlı, sıkıştırılmış, hızlı okumalı veri formatı. JSON/CSV'ye göre çok daha "
        "küçük ve hızlıdır; pandas ile native uyumludur.",
        in_project="abstracts_filtered.parquet (~1.5 GB, 633K satır) — temel veri kaynağı.")

    add_term(doc, "Veri temizleme",
        "Eksik satırları atma, NaN'leri doldurma, çok kısa abstract'ları çıkarma, "
        "duplikasyon kaldırma vs.",
        in_project="MIN_ABSTRACT_CHARS=200, MIN_TITLE_CHARS=5 kontrolleri uygulanıyor.")

    # ─── 3. Embedding ve Vektör Arama ───
    add_heading(doc, "3. Embedding ve Vektör Arama", level=1)

    add_term(doc, "ChromaDB",
        "Açık kaynak vektör veritabanı. Lokal disk üzerinde SQLite + dosya bazlı persistans, "
        "HTTP/Python client'ları, metadata filtreleme. Küçük-orta projeler için ideal.",
        in_project="633K embedding + metadata + HNSW indeks; data/chroma_db_v2/ altında 13 GB.")

    add_term(doc, "HNSW (Hierarchical Navigable Small World)",
        "Approximate Nearest Neighbor (ANN) algoritması. Milyonlarca vektör arasında log-zamanlı "
        "yakın komşu araması yapar. Tam doğru sonucu garantilemez ama çok hızlıdır.",
        in_project="Chroma içinde varsayılan; cosine space ile kullanılıyor.")

    add_term(doc, "L2 Normalizasyon",
        "Bir vektörü kendi büyüklüğüne (L2 normuna) bölerek uzunluğunu 1'e indirme. Sonrası "
        "cosine similarity = nokta çarpım haline gelir; hesap hızlanır.",
        in_project="encode(..., normalize_embeddings=True) ile her embedding L2-normalized; "
                   "13 GB indeksin doğru kosinüs vermesi buna bağlı.")

    add_term(doc, "Top-K retrieval",
        "Sorguya en yakın K dokümanı döndürme. K küçük (3-10) tutulur; daha fazlası gürültü "
        "ekler, LLM'in dikkat penceresini doldurur.",
        in_project="Retriever k=6 ile her alt-soruyu sorguluyor, dedup sonrası en iyi skorlar tutuluyor.")

    # ─── 4. Multi-Agent Mimari ───
    add_heading(doc, "4. Multi-Agent Orkestrasyonu (LangGraph)", level=1)

    add_term(doc, "Agent",
        "Belirli bir rolü olan, prompt + LLM + araç (tool) erişimine sahip bir bileşen. "
        "\"Planner\" alt sorular üretir, \"writer\" cevap yazar, \"judge\" kalite ölçer.",
        in_project="5 agent var: planner, retriever (tool), live_search, writer, critic/judge.")

    add_term(doc, "LangGraph",
        "LangChain ekibinden state machine kütüphanesi. Düğümler (node = agent) ve kenarlar "
        "(edge = condition) ile akış tanımlanır; her node bir state objesini güncelleyerek geçirir.",
        in_project="src/turk_researcher/graph.py — planner → retriever → live_search → writer → critic akışı.")

    add_term(doc, "State (durum)",
        "Tüm agent'ların okuyup yazdığı paylaşılan veri yapısı (Pydantic class). Soru, plan, "
        "chunks, synthesis, critic, final gibi alanlar tutulur.",
        in_project="State ilerledikçe doluyor; her agent kendi alanını ekliyor, sonraki okuyor.")

    add_term(doc, "Conditional edge",
        "Bir node bittikten sonra hangi node'a geçileceğini state'e bakarak belirleyen "
        "fonksiyon. \"Critic kapsama yetersiz dedi\" → \"yeni sorgularla retrieve\" gibi.",
        in_project="Critic'ten sonra 'devam et' veya 'requery yap' kararı conditional edge.")

    add_term(doc, "Tool / Function calling",
        "LLM'in dış dünyaya (API, DB, hesap) erişmesini sağlayan mekanizma. Modern API'lerde "
        "(OpenAI, Anthropic, DeepSeek) JSON şemasıyla beslenmiş 'tool definition' verilir, "
        "LLM hangi tool'u nasıl çağıracağını yapılandırılmış formda söyler.",
        in_project="Retrieval ve live search 'tool' olarak modelden bağımsız çalıştırılıyor.")

    # ─── 5. Yapılandırılmış Çıktı ───
    add_heading(doc, "5. Yapılandırılmış Çıktı (Structured Output)", level=1)

    add_term(doc, "Pydantic",
        "Python için tip kontrollü veri sınıfı kütüphanesi. dataclass benzeri ama otomatik "
        "validation yapıyor; LLM çıktılarını şemaya uydurmak için ideal.",
        in_project="Plan, Synthesis, Critic, FinalAnswer, Judgment — hepsi Pydantic BaseModel.")

    add_term(doc, "with_structured_output(method=\"function_calling\")",
        "LangChain'in LLM'i Pydantic şemasına uyacak biçimde çağırma yardımcısı. "
        "function_calling: model'in fonksiyon argümanı olarak şemayı doldurmasını ister.",
        in_project="Önce kullanıldı; DeepSeek'te sporadik None döndüğü için PydanticOutputParser'a geçildi.")

    add_term(doc, "PydanticOutputParser",
        "Daha eski/güvenilir yaklaşım: Modelden \"şu JSON şemasına uy\" prompt'uyla isteyip "
        "çıktıyı manuel parse etmek. Markdown fence (```json ... ```) varsa strip edilir.",
        in_project="space/app.py'de planner için kullanılıyor; function_calling fallback'i.")

    # ─── 6. Kontrastif Öğrenme (Stage 1 Çekirdeği) ───
    add_heading(doc, "6. Kontrastif Öğrenme — Embedder Fine-Tune'un Kalbi", level=1)

    add_para(doc,
        "Bu bölüm Stage 1'in en kritik kavramsal çekirdeğidir. Buradaki terimleri iyi "
        "kavrarsan modern representation learning'in büyük kısmı oturur."
    )

    add_term(doc, "Contrastive learning",
        "\"Benzeyen şeyleri yaklaştır, benzemeyenleri uzaklaştır\" prensibiyle representation "
        "öğrenme. Triplet learning, SimCLR, SimCSE bu ailenin parçaları.",
        in_project="trakad-embed-v2 fine-tune'unun temel paradigması.")

    add_term(doc, "Anchor / Positive / Negative",
        "Anchor referans örnek. Positive ona anlamca yakın bir örnek. Negative ona uzak bir "
        "örnek. Model anchor'ı positive'e yaklaştırmayı, negative'den uzaklaştırmayı öğrenir.",
        in_project="anchor=title_tr, positive=abstract_tr, hard_negative=aynı subject'ten farklı tezin abstract'ı.")

    add_term(doc, "InfoNCE / MultipleNegativesRankingLoss",
        "Kontrastif loss'un en yaygın hâli. Anchor'ın doğru positive'i 1 doğru cevap, batch'teki "
        "tüm diğerleri yanlış cevap olarak ele alınır. Softmax + cross-entropy ile model'in "
        "doğru cevabı seçme olasılığı maksimize edilir.",
        in_project="losses.MultipleNegativesRankingLoss(model) — batch=128 → 383 negative/anchor.")

    add_term(doc, "In-batch negatives",
        "Açıkça negative oluşturmadan, batch'teki diğer örneklerin positive'lerini negative "
        "olarak kullanma. Ücretsiz öğretmen sinyali. Batch büyüklüğü kalite kaldıracıdır.",
        in_project="Batch=128 her anchor'a 254 ekstra in-batch negative + 1 explicit hard negative verir.")

    add_term(doc, "Hard negative mining",
        "Modelin karıştırma riski yüksek olan negative'leri seçme. Random negative kolaydır; "
        "hard negative öğretici. Aynı topikten farklı dokümanlar tipik hard negative kaynağıdır.",
        in_project="Aynı subject alanından farklı tez seçilerek hard negative üretildi (subject-aware mining).")

    add_term(doc, "Catastrophic forgetting",
        "Pre-trained modeli yeni veriyle agresif eğitince eski bilgisinin silinmesi. Düşük "
        "learning rate ve az epoch ile önlenir.",
        in_project="LR=2e-5 ve 1 epoch — pre-trained mpnet bilgisini silmemek için bilinçli düşük.")

    add_term(doc, "Batch size = öğretmen kalitesi",
        "Contrastive setup'ta batch ne kadar büyükse anchor o kadar çok yanlış cevap görür, "
        "bu da o kadar keskin bir karar sınırı öğrenmesini sağlar. T4 → A100 geçişimizin ana "
        "kalite kazancı buradan.",
        in_project="A100 ile batch=128, T4'teyken yapamayacaktık.")

    # ─── 7. Eğitim Hyperparameter'leri ───
    add_heading(doc, "7. Eğitim Hyperparameter'leri (Anlamayı Öğren)", level=1)

    add_term(doc, "Learning rate (LR)",
        "Her step'te ağırlıkların ne kadar güncelleneceği. Çok yüksek = osilasyon/divergence; "
        "çok düşük = yavaş/durgun. Fine-tune için tipik 1e-5 ile 5e-5 arası.",
        in_project="LR=2e-5 — domain fine-tune için sektör standardı.")

    add_term(doc, "Warmup",
        "İlk birkaç step'te LR'i düşük başlatıp hedef değere kadar lineer artırma. Modelin "
        "ani büyük değişikliklerden korunması için.",
        in_project="Warmup ratio %3 → ilk %3 step'te LR ısınıyor.")

    add_term(doc, "Optimizer (AdamW)",
        "Adam'ın weight decay'li versiyonu; transformer eğitiminin de-facto seçimi. SGD'ye "
        "göre adaptif LR sunar.",
        in_project="Sentence-transformers fit() default'u; özelleştirilmedi.")

    add_term(doc, "Mixed precision (AMP, fp16/bf16)",
        "Eğitim sırasında bazı tensörleri fp32 yerine fp16 (yarı hassasiyet) tutarak bellek "
        "yarısını kazanma + tensor core'larda hız. Loss scaling overflow'dan korur.",
        in_project="use_amp=True; A100'de batch=128, max_seq=256 ile rahat sığdı.")

    add_term(doc, "Epoch",
        "Tüm verinin bir kez modelden geçmesi. Çoklu epoch overfitting riskini artırır; "
        "domain fine-tune için 1-3 yeter.",
        in_project="1 epoch (633K örnek, 19.812 step) — loss düzgün düştüğü için ekstraya gerek yoktu.")

    add_term(doc, "Step",
        "Bir mini-batch'in forward+backward+optimizer.step() döngüsü. 1 epoch = (örnek_sayısı / batch_size) step.",
        in_project="A100'de saniyede ~5 step; toplam 19.812 step ~68 dk.")

    add_term(doc, "Loss curve",
        "Her step'teki loss değerinin grafiği. Sağlıklı eğitimde monotonik düşer; plato/spike "
        "problemlere işaret eder.",
        in_project="0.243 → 0.062 (yumuşak düşüş, collapse yok).")

    add_term(doc, "Embedding collapse",
        "Tüm metinlerin vektörlerinin neredeyse aynı yere çökmesi. Unsupervised SimCSE, "
        "negative dropout/augmentation yetersizse bunu yapabilir. MNR loss ile in-batch + hard "
        "negative kombinasyonu çökmeyi engeller.",
        in_project="Bu projede önceki unsupervised SimCSE denemesi collapse etmişti; v2 bunu önledi.")

    # ─── 8. Mixed Precision ve GPU Bellek ───
    add_heading(doc, "8. GPU Bellek ve Performans", level=1)

    add_term(doc, "VRAM (GPU memory)",
        "Modelin ağırlıkları + gradient + optimizer state + activation'ları GPU RAM'inde tutulur. "
        "Aşılırsa OutOfMemoryError. mpnet ~280M param × 4 byte = 1.1 GB sadece ağırlık.",
        in_project="T4 (16 GB) batch=32 OOM verdi; A100 (40 GB) ile batch=128 rahat sığdı.")

    add_term(doc, "OOM (Out Of Memory)",
        "GPU bellek aşımı. Çözüm: batch küçült, max_seq düşür, gradient checkpointing aç, "
        "AMP/fp16 etkinleştir, daha büyük GPU.",
        in_project="T4 OOM → max_seq=128 + batch=16 → A100 upgrade'le geri dönüldü.")

    add_term(doc, "Tensor cores",
        "NVIDIA Volta/Turing/Ampere GPU'larındaki matris çarpımı için özel hızlandırıcı. "
        "fp16/bf16 işlemlerde 4-8x hız sağlar.",
        in_project="A100 tensor cores + AMP, T4'e göre 3-5x hızlandırdı.")

    # ─── 9. Değerlendirme ───
    add_heading(doc, "9. Değerlendirme — Sistemi Nasıl Ölçüyoruz?", level=1)

    add_term(doc, "Golden test set",
        "Manüel hazırlanmış, kategorize edilmiş soru kümesi. Her soru için beklenen subject, "
        "minimum citation sayısı vs. tanımlanmış. Sistem değişikliklerinde regresyon kontrolü için kullanılır.",
        in_project="data/eval/questions.json — 30 soru (NLP, sağlık, enerji, eğitim vs).")

    add_term(doc, "Citation accuracy",
        "Yanıttaki [n] atıfları gerçekten o kaynağı destekliyor mu? %0-100 arası.",
        in_project="LLM-as-judge bunu 0-1 skala olarak veriyor.")

    add_term(doc, "Faithfulness",
        "Yanıtın getirilen kaynaklara bağlı kalması (uydurma yok). 0-1.",
        in_project="Hallucination'ı yakalamak için ana metrik.")

    add_term(doc, "Coverage",
        "Planner'ın ürettiği alt soruların ne kadarı yanıtta gerçekten ele alınmış. 0-1.",
        in_project="Eksik konu kalıp kalmadığını ölçer.")

    add_term(doc, "Holistic score",
        "Genel akademik kalite (Türkçe akademik üslup, organizasyon, akıl yürütme). 1-5.",
        in_project="LLM-as-judge'a son cümle olarak \"genel kalite kaç?\" sordurmuş gibi.")

    add_term(doc, "LLM-as-judge",
        "Bir LLM'i (genelde başka bir model veya zero-temperature versiyon) değerlendirici "
        "olarak kullanma. İnsan değerlendirmesinin ucuz proxy'si. Calibration ve bias risklerine "
        "dikkat etmek gerek.",
        in_project="scripts/07_judge_eval.py — DeepSeek temperature=0 ile her cevabı puanlıyor.")

    # ─── 10. Türkçe-Spesifik Konular ───
    add_heading(doc, "10. Türkçe-Spesifik Konular", level=1)

    add_term(doc, "Sondan eklemeli (agglutinative) yapı",
        "Türkçe kelime köküne art arda ekler getirilerek yeni kelimeler türetilir "
        "(ev → evler → evlerimizden). Tokenizer'lar kelime yerine subword'lerle çalışınca "
        "bu yapıyı handle edebilir.",
        in_project="mpnet'in SentencePiece tokenizer'ı subword'lerle çalıştığı için Türkçeyi makul ölçüde işleyebiliyor.")

    add_term(doc, "Türkçe karakter sorunu",
        "ç, ğ, ı, İ, ö, ş, ü harfleri bazı API'lerin (OpenAlex, Semantic Scholar) sorgu "
        "parser'larını şaşırtıyor. Sonuç boş dönüyor.",
        in_project="ASCII transliteration helper + planner'a en_search_terms şeması eklendi (Türkçe→İngilizce keyword).")

    add_term(doc, "Çok dilli embedder",
        "Tek modelle 50+ dil işleyen embedder (paraphrase-multilingual-mpnet, LaBSE, vs). "
        "Tek-dilli versiyondan biraz daha düşük doğruluk verir ama dil bağımsız retrieval yapar.",
        in_project="Baseline mpnet çok dilli; v2 fine-tune Türkçe akademik domain'e specialize ediyor.")

    # ─── 11. MLOps ve Deployment ───
    add_heading(doc, "11. MLOps, Yayın ve Reproducibility", level=1)

    add_term(doc, "Hugging Face Hub",
        "Model, dataset ve Space hosting hizmeti. Git tabanlı; LFS ile büyük dosyaları handle eder. "
        "API ile programatik upload/download.",
        in_project="abstracts_filtered.parquet, trakad-embed-v2, chroma_db_v2 — hepsi Hub'da.")

    add_term(doc, "HF Spaces",
        "Streamlit/Gradio/Docker app'lerini bedava (ya da paid GPU ile) host eden servis. "
        "Repo push'la otomatik deploy.",
        in_project="hakansabunis/trakad-academic-rag-demo — Gradio 5 BYO-key demo.")

    add_term(doc, "BYO-key (Bring Your Own Key)",
        "Demo paylaşımında sunucu sahibinin API key'ini gizlemek için kullanıcının kendi key'ini "
        "girdiği pattern. UI'da password textbox; key sunucuya kayıt olmaz, sadece o session'da kullanılır.",
        in_project="DeepSeek key kullanıcıdan; quota gizlilik açısından kritik.")

    add_term(doc, "Reproducibility",
        "Aynı kodu + aynı veriyi alıp aynı sonucu elde edebilme özelliği. Seed sabitleme, "
        "version pinning, artifact versionlama, kod git history gerektirir.",
        in_project="Tüm deterministik adımlar SEED=42 ile sabit; kod GitHub'da, model+veri Hub'da.")

    add_term(doc, "Colab",
        "Google'ın bedava Jupyter notebook + GPU servisi. T4, V100, A100 (Pro+) seçenekleri. "
        "Disconnect riski var, uzun işler için Pro/Pro+ önerilir.",
        in_project="Stage 1 fine-tune ve indeks rebuild Colab A100'de yapıldı.")

    # ─── 12. Sürecin Adım-Adım Akışı ───
    add_heading(doc, "12. Adım-Adım Pipeline (Tek Sayfa)", level=1)
    add_table(doc,
        headers=["#", "Adım", "Çıktı"],
        rows=[
            ["1", "YÖK Tez OAI-PMH harvest", "raw_records.jsonl (~3 GB)"],
            ["2", "Dublin Core parse + deduplikasyon", "all_theses.parquet"],
            ["3", "Quality filter (abstract ≥200, title ≥5)", "abstracts_filtered.parquet (633K)"],
            ["4", "(v1) mpnet ile encode + Chroma index", "chroma_db/ (16 GB)"],
            ["5", "30-soru golden test + LLM-as-judge", "summary.json (v1)"],
            ["6", "Triplet üretimi (subject-aware hard neg)", "633K (anchor, positive, hard_neg)"],
            ["7", "Colab A100'de SimCSE fine-tune", "trakad-embed-v2 (1.1 GB)"],
            ["8", "(v2) trakad-embed-v2 ile encode + Chroma index", "chroma_db_v2 (13 GB)"],
            ["9", "30-soru re-eval + judge (devam ediyor)", "summary_v2.json"],
            ["10", "v1 vs v2 comparison report", "summary_diff.md"],
            ["11", "HF Hub'a push (model, indeks, parquet)", "Hub artifact'ları"],
            ["12", "HF Space'te canlı demo", "https://huggingface.co/spaces/..."],
        ],
        col_widths=[0.4, 3.2, 3.0],
    )

    # ─── 13. Öğrenmen Gerekenler ───
    add_heading(doc, "13. Öğrenmen Gereken Anahtar Kavramlar (Önem Sırasıyla)", level=1)
    add_para(doc,
        "Bu listeye sırayla bak — yukarıdan aşağı bilgi yığma sırası. Her birini "
        "anladıktan sonra bir sonrakine geç."
    )
    add_table(doc,
        headers=["Öncelik", "Kavram", "Neden önemli"],
        rows=[
            ["1", "Embedding ve cosine similarity", "Tüm retrieval'ın temeli"],
            ["2", "RAG mimarisi", "Modern LLM uygulamalarının %80'i"],
            ["3", "Transformer / BERT / GPT farkı", "Encoder vs decoder, neyi neye kullanırız"],
            ["4", "Tokenization (subword)", "Modelin metni nasıl gördüğünü anlamak"],
            ["5", "Pre-training vs fine-tuning", "Maliyet/etki dengesi"],
            ["6", "Contrastive learning + InfoNCE", "Temsil öğrenmenin ana paradigması"],
            ["7", "ChromaDB / vektör DB temel API", "Pratik retrieval"],
            ["8", "Pydantic ile yapılandırılmış çıktı", "LLM'i güvenilir bir bileşen yapar"],
            ["9", "LangGraph state machine", "Multi-agent orkestrasyonu"],
            ["10", "Mixed precision (AMP) + GPU bellek", "Eğitim ölçeklendirme"],
            ["11", "LLM-as-judge ve eval framework'leri", "Kaliteyi ölçemezsen iyileştiremezsin"],
            ["12", "HF Hub + Spaces + Colab pipeline'ı", "MLOps reproducibility"],
        ],
        col_widths=[0.7, 2.8, 3.0],
    )

    # ─── 14. Mini Sözlük ───
    add_heading(doc, "14. Mini Sözlük (Hatırlatıcı)", level=1)
    add_para(doc, "Üst bölümlerde detaylı geçen terimlerin tek satır karşılığı:", italic=True)
    add_bullet(doc, "RAG — Retrieval-Augmented Generation; LLM'i kendi corpus'unla destekleme.")
    add_bullet(doc, "Embedding — Metin → sayı vektörü dönüşümü.")
    add_bullet(doc, "Cosine similarity — İki vektörün açısının kosinüsü; benzerlik ölçüsü.")
    add_bullet(doc, "HNSW — Approximate nearest neighbor için graf algoritması.")
    add_bullet(doc, "Tokenization — Metni model girişine uygun parçalara bölme.")
    add_bullet(doc, "Pre-training — Modelin büyük genel veride sıfırdan eğitilmesi.")
    add_bullet(doc, "Fine-tuning — Pre-trained modelin küçük domain verisinde ince ayarı.")
    add_bullet(doc, "Contrastive learning — Benzeri yaklaştır, farklıyı uzaklaştır prensibi.")
    add_bullet(doc, "Anchor/positive/negative — Üçlü öğrenme şeması.")
    add_bullet(doc, "MNR loss / InfoNCE — Kontrastif öğrenmenin de-facto loss'u.")
    add_bullet(doc, "Hard negative — Modelin karıştırması muhtemel zor negative.")
    add_bullet(doc, "In-batch negatives — Batch'teki diğer pozitiflerin negative olarak kullanılması.")
    add_bullet(doc, "Catastrophic forgetting — Fine-tune'da pre-trained bilginin silinmesi.")
    add_bullet(doc, "Mixed precision (AMP) — fp16 ile bellek/hız kazancı.")
    add_bullet(doc, "Tensor core — NVIDIA GPU'larında matris çarpımı hızlandırıcı.")
    add_bullet(doc, "OOM — GPU bellek aşımı.")
    add_bullet(doc, "Citation accuracy / Faithfulness / Coverage — RAG eval metrikleri.")
    add_bullet(doc, "LLM-as-judge — LLM'i değerlendirici olarak kullanma.")
    add_bullet(doc, "Pydantic — Python'da tipli veri sınıfı + validation.")
    add_bullet(doc, "LangGraph — State machine tabanlı agent orkestrasyonu.")
    add_bullet(doc, "BYO-key — Demo'da kullanıcının kendi API key'ini girmesi pattern'i.")

    # ─── 15. İleri Okuma ───
    add_heading(doc, "15. İleri Okuma Önerileri (Kavramlar İçin)", level=1)
    add_bullet(doc, "Attention Is All You Need (Vaswani et al., 2017) — Transformer makalesi.")
    add_bullet(doc, "Sentence-BERT (Reimers & Gurevych, 2019) — Cümle embedding'inin temeli.")
    add_bullet(doc, "SimCSE (Gao et al., 2021) — Kontrastif cümle embedding fine-tune'u.")
    add_bullet(doc, "DPR / ColBERT — Akademik retrieval üzerine modern yaklaşımlar.")
    add_bullet(doc, "RAG: Lewis et al. 2020 \"Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks\".")
    add_bullet(doc, "LangChain + LangGraph dokümantasyonu — Multi-agent tasarım pattern'leri.")
    add_bullet(doc, "Hugging Face NLP Course — Ücretsiz, uygulamalı, kavramsal sağlam.")
    add_bullet(doc, "Andrej Karpathy YouTube serisi — Transformer ve LLM iç işleyişi.")

    # ─── 16. Kapanış ───
    add_heading(doc, "16. Kapanış", level=1)
    add_para(doc,
        "Bu dokümanı bir kez okuyup atılabilecek bir özet olarak değil; her bölümü "
        "kendi başına test edebileceğin bir kontrol listesi olarak düşün. Bir "
        "kavramı anladığını sınamak için en iyi yol; başkasına anlatmaya çalışmak ya da "
        "küçük bir prototip kurmaktır. TürkResearcher kodu (GitHub) tam bu amaçla şeffaf: "
        "scripts/, src/turk_researcher/, colab/ klasörlerindeki her dosya, bu rehberdeki "
        "bir kavramın somut karşılığıdır. Önce konsepti oku, sonra kodda izini sür."
    )

    doc.save(OUT)
    print(f"[+] Wrote {OUT}")
    size_kb = OUT.stat().st_size / 1024
    print(f"    {size_kb:.1f} KB")


if __name__ == "__main__":
    build()
