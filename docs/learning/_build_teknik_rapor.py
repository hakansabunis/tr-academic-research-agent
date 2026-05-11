"""Build TurkResearcher_Teknik_Rapor.docx — Stage 0 + Stage 1 retrospective.

Run:
    python docs/learning/_build_teknik_rapor.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(__file__).resolve().parent
OUT = HERE / "TurkResearcher_Teknik_Rapor.docx"


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_heading(doc, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        if level == 0:
            run.font.size = Pt(18)
        elif level == 1:
            run.font.size = Pt(13.5)
        elif level == 2:
            run.font.size = Pt(11.5)
        else:
            run.font.size = Pt(11)


def add_para(doc, text: str, *, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)


def add_meta_row(doc, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    r1.font.name = "Times New Roman"; r1.font.size = Pt(11); r1.bold = True
    r2 = p.add_run(value)
    r2.font.name = "Times New Roman"; r2.font.size = Pt(11)


def add_table(doc, headers: list[str], rows: list[list[str]],
              header_fill: str = "1F4E79", header_color=RGBColor(0xFF, 0xFF, 0xFF),
              col_widths: list[float] | None = None) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)
        run.font.color.rgb = header_color
        set_cell_shading(hdr[i], header_fill)

    for ri, row in enumerate(rows, start=1):
        cells = tbl.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)

    if col_widths:
        for row in tbl.rows:
            for c, w in zip(row.cells, col_widths):
                c.width = Inches(w)

    doc.add_paragraph()  # spacer


def build():
    doc = Document()

    # ─── Cover ───
    add_heading(doc, "TürkResearcher", level=0)
    add_heading(doc, "Teknik Retrospektif: Stage 0 + Stage 1", level=1)

    add_meta_row(doc, "Yazar", "Hakan Sabunis")
    add_meta_row(doc, "Konu", "Türkçe akademik araştırma için multi-agent RAG + custom embedder")
    add_meta_row(doc, "Kapsam", "Çok-aşamalı geliştirme süreci, teknik kararlar ve öğrenimler")
    add_meta_row(doc, "Tarih", "Mayıs 2026")
    doc.add_paragraph()

    # ─── 1. Yönetici Özeti ───
    add_heading(doc, "1. Yönetici Özeti", level=1)

    add_para(doc,
        "TürkResearcher; YÖK Tez ve DergiPark gibi parçalı Türkçe akademik kaynakların "
        "üzerinde sorgu planlayan, alıntılı sentez üreten ve kendi kalitesini ölçen bir "
        "multi-agent RAG (Retrieval-Augmented Generation) sistemidir. Proje iki ana "
        "aşamada inşa edildi:"
    )
    add_bullet(doc,
        "Stage 0 — Foundation: 633K Türkçe tezi indeksleyen Chroma + mpnet RAG; LangGraph "
        "tabanlı planner→retriever→writer→judge state machine; 30-soru golden test seti ve "
        "LLM-as-judge değerlendirme; HF Space üzerinde canlı BYO-key demo."
    )
    add_bullet(doc,
        "Stage 1 — Custom Embedder: SimCSE/MultipleNegativesRankingLoss ile fine-tune "
        "edilen trakad-embed-v2 modeli (633K tez, 1 epoch A100); subject-aware hard "
        "negative mining; eski Chroma indeksinin yeni modelle yeniden inşası ve karşılaştırmalı "
        "değerlendirme."
    )
    add_para(doc,
        "Bu rapor; her iki aşamanın teknik tasarım kararlarını, karşılaşılan engelleri, "
        "kullanılan NLP/LLM kavramlarını ve süreç boyunca öğrenilen örüntüleri özetler. "
        "Hedef; aynı problem alanında çalışacak okuyucu için bir tasarım rehberi ve "
        "kişisel bir teknik birikim kaydı sunmaktır."
    )

    # ─── 2. Stage 0 ───
    add_heading(doc, "2. Stage 0: Multi-Agent RAG (Foundation)", level=1)

    add_heading(doc, "2.1 Problem Tanımı", level=2)
    add_para(doc,
        "Türk akademik dünyasında YÖK Tez (≈1M tez), DergiPark (≈700K makale), OpenAlex "
        "ve Semantic Scholar farklı API'ler ve veri formatlarıyla erişilen birbirinden bağımsız "
        "kaynaklardır. ChatGPT/Perplexity gibi araçlar Türkçe akademik corpus'a değmemekte; "
        "alıntı üretirken hallucinate eğilimleri yüksektir. TürkResearcher'ın hedefi bu boşluğu "
        "kapatan, atıflarını gerçek kaynaklara dayandıran, açıklanabilir bir araştırma asistanı sunmaktır."
    )

    add_heading(doc, "2.2 Mimari Bileşenler", level=2)
    add_table(doc,
        headers=["Katman", "Bileşen", "Rol"],
        rows=[
            ["Veri", "OAI-PMH harvester", "YÖK Tez Dublin Core metadata akışı"],
            ["Veri", "abstracts_filtered.parquet", "633K Türkçe abstract'ın temizlenmiş hâli"],
            ["Veri", "Hugging Face Hub Dataset", "Versiyonlu, paylaşılabilir veri ve indeks deposu"],
            ["Index", "ChromaDB + HNSW (cosine)", "Yerel ve sunucusuz vektör arama"],
            ["Embedding (v1)", "paraphrase-multilingual-mpnet-base-v2", "768-dim çok dilli baseline"],
            ["LLM", "DeepSeek (OpenAI uyumlu API)", "Planner ve writer agent'larının dil modeli"],
            ["Orkestrasyon", "LangGraph state machine", "Planner→retriever→livesearch→writer→judge akışı"],
            ["Yapılandırılmış çıktı", "Pydantic + PydanticOutputParser", "JSON şema doğrulama, function_calling fallback"],
            ["Değerlendirme", "30 soru + LLM-as-judge", "Citation accuracy, faithfulness, coverage, holistic"],
            ["Deployment", "Hugging Face Space (Gradio 5)", "BYO-key canlı demo"],
        ],
        col_widths=[1.0, 2.4, 3.1],
    )

    add_heading(doc, "2.3 Önemli Stage 0 Bulguları", level=2)
    add_para(doc, "Corpus Expansion Paradox.", bold=True)
    add_para(doc,
        "Korpusu 633K tezden 740K (DergiPark'la birleşik) yapınca ortalama citation accuracy "
        "düştü. Sebep: DergiPark içeriği daha kısa, farklı dağılımdaydı; tek embedder ile iki "
        "alanı eşit kalitede temsil etmek zor olduğu için retrieval gürültüsü arttı. "
        "Bu bulgu \"daha çok veri her zaman daha iyi değildir\" prensibinin somut bir örneği "
        "olarak kayda geçti ve Stage 1 için motivasyon oluşturdu: domain-specific bir embedder "
        "eğitirsek bu gürültüyü azaltabiliriz."
    )
    add_para(doc, "Yapılandırılmış Çıktı Güvenilirliği.", bold=True)
    add_para(doc,
        "DeepSeek API'sinde with_structured_output(method=\"function_calling\") sporadik olarak "
        "None dönüyordu. Çözüm; PydanticOutputParser kullanıp model çıktısını Markdown "
        "fence'lerinden ayıklayıp manuel doğrulamak oldu. Bu deneyim API-bağımlı yapılandırılmış "
        "çıktının ne kadar kırılgan olabileceğini gösterdi ve ileride benzer sistemler için "
        "fallback parser pattern'inin değerini netleştirdi."
    )

    # ─── 3. Stage 1 ───
    add_heading(doc, "3. Stage 1: Custom Embedder Fine-Tuning", level=1)

    add_heading(doc, "3.1 Hipotez", level=2)
    add_para(doc,
        "Baseline mpnet, çok-dilli paraphrase ve STS verisinde önceden eğitilmiştir. "
        "Türkçe akademik terminolojide, jargon-yoğun cümlelerde ve uzun abstract'larda "
        "yetersiz kalmaktadır. Hipotez: title↔abstract gibi doğal pozitif çiftler ve "
        "subject-aware hard negative'lerle yapılan kontrastif fine-tune, Türkçe akademik "
        "retrieval kalitesini ölçülebilir biçimde artıracaktır."
    )

    add_heading(doc, "3.2 Kontrastif Öğrenme — Kavramsal Çerçeve", level=2)
    add_para(doc,
        "Kontrastif öğrenme; embedder'ı pozitif çiftleri yaklaştırmaya, negatifleri uzaklaştırmaya "
        "zorlayan denetimli/yarı-denetimli bir paradigmadır. Her tezde anchor (title) ve positive "
        "(abstract) kendiliğinden eşleşir; hard negative ise aynı subject alanından farklı bir "
        "tezin abstract'ı olarak seçilir. Bu seçim model'i \"yakın ama farklı\" durumda discrimination "
        "yapmaya zorlar — ki Stage 0 30-soru eval'da gözlenen başarısızlık tam olarak buydu."
    )
    add_para(doc, "Loss fonksiyonu (MultipleNegativesRankingLoss / InfoNCE):", italic=True)
    add_para(doc,
        "Her batch; B anchor + B positive + B explicit hard negative içerir. Anchor i için "
        "doğru cevap kendi positive'i, yanlış cevaplar (a) verilen explicit hard negative ve "
        "(b) batch'teki diğer anchor'ların positive'leri ve hard negative'leridir. Batch=128 "
        "iken her anchor 383 yanlış arasından doğruyu seçmeye çalışır. Batch büyüdükçe "
        "discriminative sinyal güçlenir; bu yüzden T4 yerine A100 (40 GB) tercih edildi."
    )

    add_heading(doc, "3.3 Triplet Üretimi", level=2)
    add_table(doc,
        headers=["Alan", "İçerik", "Kalite eşiği"],
        rows=[
            ["anchor", "title_tr (tez başlığı)", "≥ 5 karakter"],
            ["positive", "abstract_tr (özet)", "≥ 200 karakter"],
            ["hard_negative", "Aynı subject alanından farklı tezin abstract'ı", "Mümkün değilse rastgele"],
        ],
        col_widths=[1.5, 3.5, 1.5],
    )
    add_para(doc,
        "633,965 triplet; SEED=42 ile deterministik (yeniden üretilebilirlik için). Subject "
        "alanı listesi defaultdict(list) ile O(1) erişimle indekslendi. Aynı tezi negative "
        "olarak sürüklememek için maksimum 5 deneme ile farklı bir indeks seçildi."
    )

    add_heading(doc, "3.4 Eğitim Yapılandırması", level=2)
    add_table(doc,
        headers=["Hiperparametre", "Değer", "Gerekçe"],
        rows=[
            ["Base model", "paraphrase-multilingual-mpnet-base-v2", "768-dim, çok dilli, kanıtlanmış başlangıç"],
            ["Loss", "MultipleNegativesRankingLoss", "Açık hard neg + ücretsiz in-batch neg"],
            ["Optimizer", "AdamW", "Sentence-Transformers default"],
            ["Learning rate", "2e-5", "Pre-trained bilgiyi korumak (catastrophic forgetting'ten kaçınmak)"],
            ["Warmup", "%3", "İlk adımlarda yumuşak ısınma"],
            ["Batch size", "128 (A100 40 GB)", "Daha güçlü contrastive sinyal"],
            ["Max seq length", "256", "Ortalama abstract 100-150 token; 256 yeterli"],
            ["Mixed precision", "fp16 (use_amp=True)", "Bellek yarıya iner, A100 tensor core hızlanır"],
            ["Epochs", "1", "Domain fine-tune'da çoklu epoch overfit yapar"],
        ],
        col_widths=[1.5, 2.0, 3.0],
    )

    add_heading(doc, "3.5 Eğitim Sonuçları", level=2)
    add_para(doc,
        "Tek epoch; 19,812 step; 68.9 dakika (A100 SXM4-40GB). Loss 0.243 → 0.062 (yaklaşık 4x "
        "azalma). Eğri yumuşak; collapse (önceki unsupervised SimCSE denemesindeki gibi) "
        "gözlenmedi. Çıktı modeli hakansabunis/trakad-embed-v2 olarak HF Hub'a yayımlandı."
    )

    # ─── 4. Veri Pipeline ───
    add_heading(doc, "4. Veri Pipeline'ı ve MLOps", level=1)
    add_para(doc,
        "Reproducibility'nin merkezinde Hugging Face Hub var. Sürecin her ana çıktısı "
        "(filtered parquet, training triplets ihtiyaç olursa, fine-tune model, Chroma indeks) "
        "Hub'a versiyonlu olarak yazılıyor. Kod GitHub'da, model ve veri Hub'da, demo Space'te. "
        "Lokal makine sadece bir orkestrasyon noktası, tek nokta hatası değil."
    )
    add_table(doc,
        headers=["Aşama", "Üretim", "Tutucu (artifact store)"],
        rows=[
            ["Harvest", "raw_records.jsonl", "Lokal (≈3 GB)"],
            ["Filter", "abstracts_filtered.parquet", "HF Hub (dataset)"],
            ["Index v1", "chroma_db/", "HF Hub (dataset, alt dizin)"],
            ["Embedder v2", "trakad-embed-v2/", "HF Hub (model)"],
            ["Index v2", "chroma_db_v2/", "HF Hub (dataset, alt dizin)"],
            ["Demo", "space/app.py", "HF Space (Gradio)"],
            ["Kod", "src/, scripts/, colab/", "GitHub"],
        ],
        col_widths=[1.2, 2.5, 2.8],
    )

    # ─── 5. Index v2 İnşası ───
    add_heading(doc, "5. Yeni Chroma İndeksinin İnşası (Day 3)", level=1)
    add_para(doc,
        "Fine-tune edilen embedder, eski indeksle uyumsuzdur (vektör uzayları farklı). Bu "
        "yüzden 633,998 dokümanın tamamı trakad-embed-v2 ile yeniden encode edilip Chroma'ya "
        "yazıldı. Encode 32.6 dakika (A100, batch=256, normalize_embeddings=True), Chroma "
        "indeksleme 30.4 dakika sürdü. Önemli detay: embedding'leri Chroma'ya bir embedding_function "
        "vermek yerine doğrudan embeddings= parametresiyle geçirdik — bu, Chroma'nın tek-thread "
        "embedding fonksiyonunu çağırması yerine batch encode'un avantajını kullanmamızı sağladı "
        "(yaklaşık 10x hızlanma)."
    )
    add_para(doc,
        "Buradan doğan ikincil sorun: ChromaDB 0.5+ collection'a embedding_function "
        "verilmediğinde persisted config name=\"default\" tutuyor; sonra get_collection ile "
        "SentenceTransformerEmbeddingFunction iliştirmek istediğimizde isim çakışması "
        "ValueError'a sebep oluyor. Çözüm; lokal kodu (vectorstore.py + retriever.py) "
        "query'leri kendimiz encode edip query_embeddings= ile vermek üzere refactor etmek oldu. "
        "Bu pattern hem v1 hem v2 koleksiyonlarıyla sorunsuz çalışıyor."
    )

    # ─── 6. Karşılaşılan Engeller ───
    add_heading(doc, "6. Karşılaşılan Engeller ve Çözümler", level=1)
    add_table(doc,
        headers=["Sorun", "Belirti", "Çözüm"],
        rows=[
            [
                "T4 OOM",
                "smoke fit() backward'da CUDA OOM",
                "max_seq=128 + use_amp=True; ardından A100 upgrade ile batch=128 ve max_seq=256 geri",
            ],
            [
                "Chroma 0.5 ef çakışması",
                "get_collection \"default vs sentence_transformer\" hatası",
                "embedding_function bağlamadan aç, query'leri kendimiz encode et",
            ],
            [
                "function_calling None",
                "DeepSeek planner çıktısı sporadik None",
                "PydanticOutputParser ile fallback; JSON fence'leri strip",
            ],
            [
                "Türkçe karakter API'leri kırması",
                "OpenAlex/Semantic Scholar boş döndürüyor",
                "ASCII transliteration helper + planner'a en_search_terms şeması",
            ],
            [
                "HF Space audioop ImportError",
                "Python 3.13 stdlib kaldırdı",
                "audioop-lts paketi requirements'a eklendi",
            ],
            [
                "API key sızıntısı endişesi",
                "Demo'yu açık paylaşmak isteniyor ama key korumak gerek",
                "BYO-key (Bring Your Own Key) mimarisi: kullanıcı key'ini UI'a girer, sunucuya gitmez",
            ],
        ],
        col_widths=[1.7, 2.5, 2.3],
    )

    # ─── 7. NLP + LLM Kavşağı ───
    add_heading(doc, "7. NLP ve LLM Tarafları (Disiplin Haritası)", level=1)
    add_para(doc,
        "Proje hem klasik NLP hem modern LLM uygulamasını kapsar. Bu birleşim modern üretim "
        "sistemlerinin tipik stack'idir; iki disiplinin nerede ayrılıp nerede birleştiğini "
        "açıkça görmek mühendislik olgunluğu için önemlidir."
    )
    add_table(doc,
        headers=["Disiplin", "Bu projedeki bileşenler"],
        rows=[
            ["Klasik NLP", "Tokenization, sentencepiece, sequence length yönetimi, transliteration, OAI-PMH/Dublin Core parsing, abstract filtering"],
            ["Information Retrieval", "ChromaDB + HNSW + cosine, doc concat formatı, query/doc embedding tutarlılığı"],
            ["Representation Learning", "SimCSE/contrastive learning, MultipleNegativesRankingLoss, hard negative mining, batch size = öğretmen kalitesi"],
            ["LLM Engineering", "DeepSeek API, structured output (Pydantic), prompt mühendisliği, planner→writer ayrımı"],
            ["Agent Orchestration", "LangGraph state machine, conditional edges, retry logic, citation discipline"],
            ["Evaluation", "30-soru golden test, LLM-as-judge (citation accuracy, faithfulness, coverage, holistic), v1↔v2 karşılaştırması"],
            ["MLOps", "HF Hub artifact store, HF Space deploy, Colab T4/A100 training, GitHub-merkezli reproducibility"],
        ],
        col_widths=[2.0, 4.5],
    )

    # ─── 8. Genel Kazanımlar ───
    add_heading(doc, "8. Süreç Boyunca Edinilen Teknik Kazanımlar", level=1)
    add_bullet(doc,
        "Multi-agent state machine'in karmaşık iş akışlarında modüler hata izolasyonu "
        "sağladığı; her node'un kendi başına test edilebilir olduğu."
    )
    add_bullet(doc,
        "Embedding kalitesinin retrieval kalitesini doğrudan belirlediği; LLM ne kadar iyi "
        "olursa olsun yanlış chunk'larla çıktı kalitesinin tavanlandığı."
    )
    add_bullet(doc,
        "Hard negative mining'in, in-batch negatives'a kıyasla daha keskin discriminative "
        "sinyal verdiği; özellikle aynı topikten gelen yakın örnekler için."
    )
    add_bullet(doc,
        "Domain fine-tune'un düşük learning rate (2e-5) gerektirdiği; pre-trained "
        "ağırlıkların \"silinmesi\" değil, \"kaydırılması\" hedeflenmesi gerektiği."
    )
    add_bullet(doc,
        "Mixed precision (fp16) eğitiminin sadece bellek değil, A100 tensor core'larında "
        "ciddi zaman kazandırdığı; T4'ten A100'e geçişin ölçek ve kalite açısından kritik kaldıraç olduğu."
    )
    add_bullet(doc,
        "ChromaDB gibi yarı-olgun araçlarda persisted config'in (embedding function, distance metric) "
        "kütüphane sürümleri arasında uyumsuzluk yaratabildiği; uçtan uca (encode + query) "
        "tek bir tutarlı yolu tercih etmenin daha sağlam olduğu."
    )
    add_bullet(doc,
        "BYO-key deployment pattern'inin demo paylaşımı için cost ve gizlilik açısından "
        "optimal olduğu; özellikle ders/portfolyo projeleri için kritik."
    )
    add_bullet(doc,
        "LLM-as-judge'un evaluation maliyetini düşürdüğü ama kalibrasyona dikkat gerektirdiği; "
        "judge prompt'una tüm chunk'ları görünür kılmak gibi küçük detayların metric güvenirliğini büyük ölçüde değiştirdiği."
    )

    # ─── 9. Sonraki Adımlar ───
    add_heading(doc, "9. Sonraki Adımlar (Stage 1 Day 4-5 ve Sonrası)", level=1)
    add_bullet(doc,
        "Day 4: 30-soru re-eval (v2 indeks + trakad-embed-v2) ve LLM-as-judge skorlarının "
        "alınması — pre/post quantitative tablo bu çıktıdan üretilecek."
    )
    add_bullet(doc,
        "Day 5: v1↔v2 karşılaştırma raporu (citation accuracy, faithfulness, coverage, "
        "holistic delta'ları) ve teknik blog post."
    )
    add_bullet(doc,
        "Stage 2 (planlanan): SFT — küçük bir base LLM (Llama-3-8B veya Qwen2-7B) üzerinde "
        "Türkçe akademik writer'ı supervised fine-tune; capstone yanıt formatlarını instruction "
        "olarak kullanma."
    )
    add_bullet(doc,
        "Stage 3 (planlanan): DPO/RLHF — judge skorlarını preference verisi olarak kullanıp "
        "writer'ı alıntı disiplinine göre alignment."
    )

    # ─── Kapanış ───
    add_heading(doc, "10. Kapanış", level=1)
    add_para(doc,
        "TürkResearcher; üniversite final projesi formatının ötesinde, üretim sistemlerinin "
        "küçük ama tam bir prototipi olarak inşa edildi: gerçek veri, fine-tune edilmiş model, "
        "canlı demo, ölçülebilir kalite. Süreç boyunca edinilen kararlar (embedder seçiminden "
        "structured output mimarisine, deployment modelinden eval framework'üne kadar) hem "
        "akademik bir teslim için hem de gelecekteki benzer NLP+LLM projeleri için tekrar "
        "kullanılabilir bir kalıp oluşturuyor. Stage 1 sonu bu kalıbın \"baseline + custom "
        "model + ölçülebilir kazanım\" döngüsünün ilk tam turunu temsil ediyor; Stage 2-3 "
        "aynı disiplinin LLM jeneratif tarafına taşınmasıdır."
    )

    doc.save(OUT)
    print(f"[+] Wrote {OUT}")
    size_kb = OUT.stat().st_size / 1024
    print(f"    {size_kb:.1f} KB")


if __name__ == "__main__":
    build()
