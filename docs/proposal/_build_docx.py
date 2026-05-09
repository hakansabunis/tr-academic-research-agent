"""Build proposal.docx — visual version with embedded block diagrams,
tables, and bullets. No code blocks, no wall-of-text.

Run:
    python docs/proposal/_build_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(__file__).resolve().parent
OUT = HERE / "TurkResearcher_Proposal_Hakan_Sabunis.docx"
ARCH_PNG = HERE / "fig_architecture.png"
DATA_PNG = HERE / "fig_data_pipeline.png"


# ─── helpers ───


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_heading(doc, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        if level == 0:
            run.font.size = Pt(18)
        elif level == 1:
            run.font.size = Pt(13.5)
        elif level == 2:
            run.font.size = Pt(11.5)
        else:
            run.font.size = Pt(11)


def add_para(doc, text: str, *, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)


def add_meta_row(doc, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    r1.font.name = "Times New Roman"; r1.font.size = Pt(11); r1.bold = True
    r2 = p.add_run(value)
    r2.font.name = "Times New Roman"; r2.font.size = Pt(11)


def add_numbered(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)


def add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[i], "1F3A5F")

    for r_idx, row_data in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = ""
            run = cells[c_idx].paragraphs[0].add_run(val)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)


def add_image(doc, path: Path, *, width_inches: float = 6.4) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))


def add_caption(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = RGBColor(0x5A, 0x66, 0x79)


# ─── content ───


def main() -> None:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Project Proposal — TürkResearcher")
    run.font.name = "Times New Roman"; run.font.size = Pt(18); run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("A Multi-Agent LLM Research Assistant for Turkish Academic Literature")
    run.font.name = "Times New Roman"; run.font.size = Pt(13); run.italic = True

    doc.add_paragraph()

    add_meta_row(doc, "Student", "Hakan Sabuniş")
    add_meta_row(doc, "Course", "Large Language Models — Spring 2026")
    add_meta_row(doc, "Institution", "Istanbul Medipol University")
    add_meta_row(doc, "Track", "Track 1 — Novel Idea")
    add_meta_row(doc, "Date", "9 May 2026")
    doc.add_paragraph()

    # ── 1. Problem ──
    add_heading(doc, "1. Problem & Motivation", 1)
    add_bullet(doc,
        "English-speaking researchers have grounded LLM tools (Elicit, Consensus.app, "
        "scite.ai); the Turkish academic ecosystem has none.")
    add_bullet(doc,
        "Turkey publishes 600+ active peer-reviewed journals on DergiPark and tens of "
        "thousands of theses on YÖK Ulusal Tez Merkezi every year.")
    add_bullet(doc,
        "All of that Turkish-language scholarly output is invisible to current "
        "LLM-based research assistants. TürkResearcher fills exactly this gap.")

    # ── 2. Architecture ──
    add_heading(doc, "2. System Architecture", 1)
    add_image(doc, ARCH_PNG, width_inches=6.6)
    add_caption(doc, "Figure 1. LangGraph state machine. Solid arrows are deterministic; "
                     "dashed arrows are conditional.")
    add_bullet(doc,
        "Six-node directed graph orchestrated by LangGraph; typed Pydantic state; "
        "five LLM-using agents plus one optional live-search node.")
    add_bullet(doc,
        "Critic conditionally loops back to Retriever (max 2 iterations), then falls "
        "through to LiveSearch or directly to Writer.")
    add_bullet(doc,
        "Every LLM node uses schema-enforced output through OpenAI function calling — "
        "DeepSeek does not support the newer json_schema response format.")

    # ── 3. Component Specifications ──
    add_heading(doc, "3. Component Specifications", 1)

    add_heading(doc, "3.1 Embedding layer", 2)
    add_table(doc,
        ["Property", "Value", "Rationale"],
        [
            ["Model", "paraphrase-multilingual-mpnet-base-v2",
             "768-dim multilingual sentence transformer; XLM-RoBERTa backbone with strong Turkish coverage"],
            ["Document representation", "Title concatenated with abstract",
             "Title-aware indexing adds a strong lexical signal that abstract-only embeddings miss"],
            ["Normalisation", "L2", "Required for cosine similarity to be a valid metric"],
            ["Inference", "CPU at query time", "Local agent runs without GPU; ~50 ms/query"],
        ])
    add_para(doc,
        "Alternatives considered and rejected: BERTurk-base (NER-tuned, weak as sentence "
        "encoder), distiluse-multilingual (smaller but lower retrieval quality), custom "
        "SimCSE fine-tune (deferred to future work).", italic=True, size=10)

    add_heading(doc, "3.2 Vector indexing", 2)
    add_table(doc,
        ["Property", "Value", "Rationale"],
        [
            ["Backend", "ChromaDB persistent client",
             "Python-native, persistent on disk, mature LangChain integration"],
            ["Index", "HNSW (M=16, ef_construction=200)",
             "Sub-second similarity search on 10⁵+ vectors"],
            ["Distance", "Cosine", "Standard for L2-normalised sentence embeddings"],
            ["Identifier", "Thesis: YÖK ID; article: prefixed",
             "Globally unique; supports source-aware filtering"],
            ["Metadata fields", "13 (author, advisor, location, year, subject, PDF URL, …)",
             "Sufficient to construct IEEE citations without re-querying parquet"],
        ])

    add_heading(doc, "3.3 LLM integration layer", 2)
    add_table(doc,
        ["Property", "Value"],
        [
            ["Model", "DeepSeek deepseek-chat"],
            ["Transport", "OpenAI-compatible REST (LangChain ChatOpenAI with overridden base URL)"],
            ["Structured output", "function_calling (json_schema unsupported by DeepSeek)"],
            ["Per-agent temperature", "Critic 0.0 · Planner 0.1 · Synthesiser 0.2 · Writer 0.3"],
        ])

    add_heading(doc, "3.4 Retrieval strategy", 2)
    add_bullet(doc,
        "Multi-query: original question plus each planner sub-question (4–6 queries per turn).")
    add_bullet(doc,
        "Per query, top-k=6 chunks by cosine similarity; deduplicated globally on thesis identifier "
        "with max-score retention; truncated to top-24 for the synthesiser.")
    add_bullet(doc,
        "ChromaDB cosine distance is converted to a [0, 1] similarity score for downstream "
        "interpretability.")

    add_heading(doc, "3.5 Citation subsystem", 2)
    add_bullet(doc,
        "The writer emits inline numeric markers; the IEEE reference list is built "
        "deterministically by the system from chunk metadata, not by the LLM.")
    add_bullet(doc,
        "Each citation includes the public YÖK detail page URL "
        "(tez.yok.gov.tr/UlusalTezMerkezi/...), where the abstract, full metadata, and "
        "the PDF (for open-access theses) are retrievable.")
    add_bullet(doc,
        "A safety pass strips any inline marker that exceeds the available chunk count — "
        "no hallucinated references can leak into the final answer.")

    # ── 4. Data Pipeline ──
    add_heading(doc, "4. Data Pipeline", 1)
    add_image(doc, DATA_PNG, width_inches=6.6)
    add_caption(doc, "Figure 2. Data pipeline from Hugging Face Hub upstream to local agent runtime.")
    add_bullet(doc,
        "Source: umutertugrul/turkish-academic-theses-dataset (650K bilingual abstracts, "
        "CC-BY-4.0).")
    add_bullet(doc,
        "Filters: abstract ≥ 50 words, non-empty Turkish title, deduplication on thesis "
        "identifier — yields ~500-600K records.")
    add_bullet(doc,
        "The prebuilt index is published as a public Hugging Face dataset so the runtime "
        "can be reproduced in three commands without rebuilding the index.")

    # ── 5. Evaluation ──
    add_heading(doc, "5. Evaluation Framework", 1)
    add_bullet(doc,
        "30-question Turkish benchmark spanning 10 categories: computer science, "
        "education, health, engineering, social sciences, business, agriculture, "
        "linguistics, law, multi-domain.")
    add_table(doc,
        ["Metric", "Range", "Captures"],
        [
            ["Citation accuracy", "0–1", "Whether cited chunks support the surrounding claim"],
            ["Faithfulness", "0–1", "Whether every claim is grounded in retrieved context"],
            ["Coverage", "0–1", "Fraction of planner sub-questions substantively addressed"],
            ["Holistic", "1–5", "Overall academic quality (structure, register, reasoning)"],
        ])
    add_bullet(doc,
        "Mechanical metrics also captured: end-to-end latency, retrieved-chunk count, "
        "max cosine similarity, critic iteration count.")
    add_bullet(doc,
        "Per-category breakdown surfaces domains where the corpus is strong vs weak — "
        "honest error analysis instead of a blended single number.")

    # ── 6. Reproducibility ──
    add_heading(doc, "6. Reproducibility & Deployment", 1)
    add_bullet(doc, "Code released under MIT licence on GitHub.")
    add_bullet(doc, "Prebuilt Chroma index published as a public Hugging Face dataset.")
    add_bullet(doc,
        "Three-step runtime: install pinned requirements, pull the prebuilt index, run "
        "the agent CLI with a Turkish question.")

    # ── 7. Timeline ──
    add_heading(doc, "7. Timeline", 1)
    add_table(doc,
        ["Phase", "Weeks", "Output"],
        [
            ["Data pipeline (fetch, filter, embed, index, push)", "1–2",
             "Indexed Chroma collection on HF Hub"],
            ["Multi-agent system implementation", "3", "Runnable agent CLI"],
            ["Evaluation framework (benchmark, judge, summary)", "4", "Quantitative results table"],
            ["Academic report (IEEE) and class presentation", "5", "PDF and slide deck"],
        ])

    # ── 8. Deliverables ──
    add_heading(doc, "8. Expected Deliverables", 1)
    for txt in [
        "GitHub repository under MIT licence with reproducible pipeline and full evaluation harness.",
        "Public Hugging Face dataset hosting the prebuilt vector index.",
        "IEEE-format academic report (2–6 pages): abstract, introduction, related work, methodology, experiments, discussion / error analysis, conclusion.",
        "10–15 minute class presentation with a live demo.",
    ]:
        add_numbered(doc, txt)

    doc.save(OUT)
    print(f"Wrote: {OUT}")
    print(f"Size: {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
